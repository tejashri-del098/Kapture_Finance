from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "Kapture_Collections_Voicebot_HLD.docx"
OUTPUT = ROOT / "Kapture_Collections_Voicebot_HLD_Corrected.docx"

REPLACEMENTS = {
    "only “this is Kapture Finance regarding your account, please call us back on …”": (
        "only “this is Maya calling from Kapture Finance; please return our call using the official contact details”"
    ),
    "the purpose (“calling about your loan account”)": (
        "the neutral purpose (“calling regarding a confidential customer-service matter”)"
    ),
    "(07:00–19:00 local time)": "(08:00–19:00 local time)",
    "Household members who correctly answer the DOB/ID challenge are, by design, treated as authorized — this mirrors how the client's existing IVR/agent process authenticates, and is called out as a residual risk in §9 assumptions.": (
        "For this demo, DOB plus ID-last-four is the mocked verification control. Production use requires the lender's compliance-approved identity-verification method (for example, an OTP or a stronger customer-authentication control); knowledge of these details alone must not be treated as third-party authorization."
    ),
    "A response filter runs on every LLM output before TTS: while auth_verified is false, it strips/blocks any output containing amount patterns, the word “EMI/loan/overdue” combined with a number, or the account holder’s full name plus financial terms. This is defense-in-depth on top of the prompt instruction, so a jailbreak or clever rephrasing by the caller cannot leak debt information even if the model is momentarily talked into attempting it.": (
        "For this Vapi demo, debt facts are never placed in the pre-auth model context and the webhook blocks every account-data or payment tool until auth_verified is set by verify_customer. A production deployment additionally requires an output-policy filter before TTS to block any pre-auth debt disclosure attempt."
    ),
    "Two-factor DOB + last-4-ID verification is assumed sufficient per the client’s existing IVR/agent standard; stronger verification (OTP-based) is flagged as a future improvement, not built here.": (
        "DOB plus last-4-ID is a mock-only exercise control. Production must use a lender-approved authentication method; this build does not authorize third parties based on knowledge of these details."
    )
}


def replace_paragraph_text(paragraph, old, new):
    if old not in paragraph.text:
        return False
    # These four target paragraphs use uniform body/bullet text. Replacing their
    # text preserves the paragraph-level style, numbering, indentation and layout.
    paragraph.text = paragraph.text.replace(old, new)
    return True


document = Document(SOURCE)
counts = {old: 0 for old in REPLACEMENTS}

for paragraph in document.paragraphs:
    for old, new in REPLACEMENTS.items():
        if replace_paragraph_text(paragraph, old, new):
            counts[old] += 1

for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for old, new in REPLACEMENTS.items():
                    if replace_paragraph_text(paragraph, old, new):
                        counts[old] += 1

missing = [text for text, count in counts.items() if count != 1]
if missing:
    raise RuntimeError(f"Expected each HLD correction once; counts: {counts}")

document.save(OUTPUT)
print(OUTPUT)
